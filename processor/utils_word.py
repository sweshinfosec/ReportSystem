import os
import re
import time
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from django.conf import settings


# --- 1. HELPER FUNCTIONS ---

def set_cell_background(cell, fill):
    """Sets cell background color via HEX."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_shading(cell, color):
    """Helper to set cell background color (e.g., '1F497D' for Blue)."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def zero_table_indent(table):
    """Explicitly pins a table's left indent to 0 so it starts flush with
    body text/headings above it, regardless of any inherited table-style
    indent default."""
    tblPr = table._tbl.tblPr
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '0')
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

def set_cell_borders_all(cell, sz=4, color="000000"):
    """Explicitly stamps all 4 borders on a cell with identical values.
    Cells that were merged from several originally-separate cells can end
    up resolving their borders ambiguously (inherited from the table style
    vs. whatever the pre-merge cells had), which some renderers show as a
    faint doubled/ghost line at the seam. Forcing explicit, identical
    borders removes that ambiguity."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), str(sz))
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), color)
        borders.append(node)
    tcPr.append(borders)

def set_cell_margins(cell, top=120, bottom=120, left=120, right=120):
    """Adds internal cell padding (in twentieths of a point / dxa) so text
    isn't cramped against the borders. Word's default is ~0 top/bottom."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def pad_header_cell(cell, top=90, bottom=90):
    """Adds breathing room + vertical centering to a colored header/label
    cell, so its text isn't glued to the top border (used on every navy
    header bar and table header row in the document)."""
    set_cell_margins(cell, top=top, bottom=bottom)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_row_height(row, height_inches, rule='atLeast'):
    """Forces a table row to a specific height instead of auto-fitting to
    content, so a set of rows line up at a consistent height. w:trHeight's
    w:val is in dxa (twentieths of a point = 1/1440 inch), NOT EMU, so it
    must be computed directly rather than via int(Inches(...))."""
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_inches * 1440)))
    trHeight.set(qn('w:hRule'), rule)
    trPr.append(trHeight)

def clean_wrapped_text(text):
    """Normalizes hard-wrapped plain-text fields (common in Nessus CSV
    exports, where a sentence is hand-wrapped with a line break every ~80
    characters). Collapses those mid-sentence line breaks into spaces so
    the paragraph flows and justifies normally, while still treating a
    genuine blank line (a real paragraph break in the source) as a break."""
    if text is None:
        return ""
    text = str(text)
    if text.strip().lower() == 'nan':
        return ""
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    blocks = re.split(r'\n\s*\n', text)
    cleaned = [re.sub(r'\s*\n\s*', ' ', block).strip() for block in blocks]
    return '\n'.join(b for b in cleaned if b)

def add_cleaned_text_block(cell, raw_text, font_size=10, fallback="N/A"):
    """Renders a narrative field (Description/Synopsis/Recommendation) into
    a cell, one real paragraph per preserved block. Justify alignment only
    ever leaves the LAST line of a paragraph unstretched - if multiple
    blocks were joined into a single paragraph via soft line breaks, every
    short block except the final one gets stretched into huge word gaps.
    Giving each block its own paragraph means each one's own last (and
    often only) line is exempt, so short lines/bullets no longer balloon."""
    cleaned = clean_wrapped_text(raw_text)
    blocks = [b for b in cleaned.split('\n')] if cleaned else []
    if not blocks:
        blocks = [fallback]
    for idx, block in enumerate(blocks):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        apply_styling(p, block, font_size=font_size)
        p.paragraph_format.space_after = Pt(6)
    return cell

def apply_styling(paragraph, text="", is_bold=False, font_size=10, color=None):
    """Applies Calibri styling and strictly removes paragraph spacing."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    
    run = paragraph.add_run(text)
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run.font.size = Pt(font_size)
    run.bold = is_bold
    if color:
        run.font.color.rgb = color
    return run

def add_page_number_field(run):
    """Adds the dynamic PAGE field to a run."""
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = "PAGE"
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2)

def add_numpages_field(run):
    """Adds the dynamic NUMPAGES field to a run (total page count)."""
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = "NUMPAGES"
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2)

def set_paragraph_border(paragraph, top=False, bottom=False, color="7F7F7F", size=6):
    """Adds a thin horizontal rule above/below a paragraph (the running
    header/footer rule lines used in the reference report)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    for side, enabled in (('top', top), ('bottom', bottom)):
        if enabled:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:val'), 'single')
            node.set(qn('w:sz'), str(size))
            node.set(qn('w:space'), '4')
            node.set(qn('w:color'), color)
            p_bdr.append(node)
    p_pr.append(p_bdr)


    # Define the blue color from your screenshot
SECTION_BLUE = RGBColor(54, 95, 145)
NAVY_BLUE = RGBColor(0x1F, 0x49, 0x7D)  # Standard navy used across headers/tables

PROJECT_PLACEHOLDER = '<Project>'

def get_project_name(workspace):
    """Returns the project name, defaulting to the unfilled placeholder
    token when the caller hasn't supplied one yet."""
    return workspace.get('project_name', PROJECT_PLACEHOLDER)

def add_text_with_project_name(paragraph, before, project_name, after, font_size=10, is_bold=False, color=None):
    """Adds `before + project_name + after` as normal text, but if
    project_name is still the unfilled placeholder, highlights just that
    portion in yellow so it's easy to spot and swap out before the report
    goes final."""
    if before:
        apply_styling(paragraph, before, is_bold=is_bold, font_size=font_size, color=color)
    val_run = apply_styling(paragraph, project_name, is_bold=is_bold, font_size=font_size, color=color)
    if project_name == PROJECT_PLACEHOLDER:
        val_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if after:
        apply_styling(paragraph, after, is_bold=is_bold, font_size=font_size, color=color)

CUSTOMER_PLACEHOLDER = '<customer>'

def get_customer_name(workspace):
    """Returns the customer/client team name (previously hardcoded as
    'NTT'/'NTTD' throughout the document), defaulting to the unfilled
    placeholder token when the caller hasn't supplied one yet."""
    return workspace.get('customer_name', CUSTOMER_PLACEHOLDER)

def add_text_with_customer_name(paragraph, before, customer_name, after, font_size=10, is_bold=False, color=None):
    """Same idea as add_text_with_project_name, but for the customer/client
    team name."""
    if before:
        apply_styling(paragraph, before, is_bold=is_bold, font_size=font_size, color=color)
    val_run = apply_styling(paragraph, customer_name, is_bold=is_bold, font_size=font_size, color=color)
    if customer_name == CUSTOMER_PLACEHOLDER:
        val_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if after:
        apply_styling(paragraph, after, is_bold=is_bold, font_size=font_size, color=color)

def _strip_extra_paragraphs(cell):
    """Merging table cells leaves one leftover empty paragraph behind from
    every cell that got merged away, which shows up as extra blank lines
    (an extra navy line under header text, or a gap before/after content).
    Keep only the first paragraph."""
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)

def format_full_width_header(table, row_idx, text):
    """Merges cells and formats as dark blue header."""
    row = table.rows[row_idx]
    # Adjust merge if your table has more/less than 3 columns
    merged_cell = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    _strip_extra_paragraphs(merged_cell)
    set_cell_background(merged_cell, "1F497D")
    set_cell_borders_all(merged_cell)
    # Table headers usually stay white for contrast
    apply_styling(merged_cell.paragraphs[0], text, is_bold=True, font_size=10, color=RGBColor(255, 255, 255))
    merged_cell.paragraphs[0].alignment = 0 # Left
    pad_header_cell(merged_cell)
    return merged_cell

def format_full_width_value(table, row_idx):
    """Merges cells for data content. Gets the same top/bottom breathing
    room as the header bars - without it, a single line of text (e.g. the
    Vulnerability name or Recommendation) sits flush against both borders
    and looks glued to the header bar above and the one below it."""
    row = table.rows[row_idx]
    merged_cell = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    _strip_extra_paragraphs(merged_cell)
    set_cell_margins(merged_cell, top=90, bottom=90)
    return merged_cell

def style_section_heading(paragraph, size=None):
    """Forces numbered section headings (1, 2, 2.1, 3.1, 4...) to plain
    black Calibri, matching the reference report. Front-matter card titles
    (SECURITY TEAM DETAILS, etc.) are left untouched and keep their navy color."""
    for run in paragraph.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
        if size:
            run.font.size = Pt(size)
    # Breathing room above/below the heading itself (~1.5 lines combined
    # with the body paragraph's own space_after below).
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(10)
    return paragraph

def format_unique_servers(rows):
    """Formats 'Host(Port:X)' entries for a 'Servers: ...' list, collapsing
    exact duplicate Host+Port pairs down to a single mention (order
    preserved). A host repeated with a DIFFERENT port is kept as a
    separate entry, e.g. 10.18.100.1(Port:445) and 10.18.100.1(Port:443)."""
    seen = set()
    parts = []
    for r in rows:
        key = (getattr(r, 'Host', None), getattr(r, 'Port', None))
        if key in seen:
            continue
        seen.add(key)
        parts.append(f"{r.Host}(Port:{r.Port})")
    return ", ".join(parts)

def compute_overall_status(status_values):
    """A single finding (e.g. 'ASP.NET Core SEoL') can appear on many
    servers, each with its own current status. The finding as a whole
    should only read as 'Fixed & Closed' once EVERY affected server has
    been fixed - if even one server out of a hundred is still 'Open',
    the finding is still 'Open'."""
    return 'Open' if 'Open' in list(status_values) else 'Fixed & Closed'

def add_body_paragraph(doc, text, font_size=10):
    """Adds a justified narrative paragraph directly under a numbered
    heading. Left-indent removed so the paragraph's left edge lines up
    exactly with the heading text and any table that follows it, instead
    of zig-zagging in from the margin."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    apply_styling(p, text, font_size=font_size)
    # Gap before whatever comes next (next heading or table), roughly a
    # blank line, so content doesn't run straight into the next section.
    p.paragraph_format.space_after = Pt(14)
    return p

def add_body_paragraph_with_project_name(doc, before, project_name, after, font_size=10):
    """Same layout/spacing as add_body_paragraph, but splits the text
    around the project name so it can be highlighted individually if it's
    still the unfilled placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text_with_project_name(p, before, project_name, after, font_size=font_size)
    p.paragraph_format.space_after = Pt(14)
    return p

# --- 2. MAIN GENERATOR ---

def generate_final_word_report(df, workspace=None):
    if workspace is None:
        workspace = {}
    
    doc = Document()

    element_update = OxmlElement('w:updateFields')
    element_update.set(qn('w:val'), 'true')
    doc.settings.element.append(element_update)
    
    # Global Style Setup
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # --- COVER PAGE (Page 1): logo, title block, CREST badge + footer at the bottom ---
    for _ in range(6): doc.add_paragraph(" ")

    logo_path = os.path.join(settings.BASE_DIR, 'static', 'images', 'logo.png')
    if os.path.exists(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.add_run().add_picture(logo_path, width=Inches(3.5))

    for _ in range(5): doc.add_paragraph(" ")

    p_project = doc.add_paragraph()
    p_project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_with_project_name(p_project, '', get_project_name(workspace), '', is_bold=True, font_size=24, color=NAVY_BLUE)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_styling(p_title, '\nVulnerability Assessment Review\nFinal Report 2026', is_bold=True, font_size=18, color=NAVY_BLUE)

    p_disclaimer = doc.add_paragraph()
    p_disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_styling(p_disclaimer, '\nCONFIDENTIAL DOCUMENT', is_bold=True, font_size=12)

    cert_path = os.path.join(settings.BASE_DIR, 'static', 'images', 'cert.png')

    # Push the CREST badge + accreditation line further down so it (and the
    # copyright block right after it) reads as the last thing on the page,
    # instead of floating mid-page with a big empty gap below it.
    for _ in range(10): doc.add_paragraph(" ")

    if os.path.exists(cert_path):
        cert_para = doc.add_paragraph()
        cert_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cert_para.add_run().add_picture(cert_path, width=Inches(1.5))

    crest_text_p = doc.add_paragraph()
    crest_text_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_styling(crest_text_p, 'We are a CREST accredited organization for Vulnerability Assessment and Penetration Testing', is_bold=False, font_size=10)

    # Small gap (not the old large one) so the copyright block sits close
    # behind the CREST line as one bottom-anchored group.
    for _ in range(2): doc.add_paragraph(" ")

    for line in ['CONFIDENTIAL DOCUMENT', 'Copyright © 2026 Opensource.', 'All Rights Reserved.']:
        copy_p = doc.add_paragraph()
        copy_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_styling(copy_p, line, is_bold=False, font_size=8)

    doc.add_page_break()

    # --- REPORT BODY (Page 2 Onwards) ---
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.header.is_linked_to_previous = False
    new_section.footer.is_linked_to_previous = False

    # Force 1-inch margins and calculate right corner (8.5" - 2" = 6.5")
    new_section.left_margin = Inches(1.0)
    new_section.right_margin = Inches(1.0)
    RIGHT_CORNER = Inches(6.5)

    # --- HEADER: Title (Left, bold navy italic) | Logo (Right) | rule line below ---
    # Uses an independent 2-column table (like the footer) instead of a tab
    # stop, so the logo always sits at the true right margin regardless of
    # how long the project name text is - the two sides never affect each other.
    header_table = new_section.header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(5.0)
    header_table.columns[1].width = Inches(1.5)

    h_left_cell = header_table.cell(0, 0)
    h_right_cell = header_table.cell(0, 1)

    h_left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    h_left_para = h_left_cell.paragraphs[0]
    h_left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_h_left = apply_styling(h_left_para, str(workspace.get('project_name', 'Security Assessment Report')).upper(), is_bold=True, font_size=10, color=NAVY_BLUE)
    run_h_left.italic = True

    h_right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    h_right_para = h_right_cell.paragraphs[0]
    h_right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(logo_path):
        h_right_para.add_run().add_picture(logo_path, width=Inches(1.0))

    for cell in (h_left_cell, h_right_cell):
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '7F7F7F')
        borders.append(bottom)
        tcPr.append(borders)

    # A fresh section header starts with one empty default paragraph, which
    # add_table() leaves sitting BEFORE the table. Even with zero spacing,
    # that empty line still takes up its own line height, which pushes the
    # table (and the header text inside it) down toward the bottom of the
    # fixed header area instead of it sitting centered. Remove that leading
    # empty paragraph outright so the table is the first thing in the header.
    hdr_body = new_section.header._element
    if len(hdr_body) and hdr_body[0].tag.endswith('}p'):
        first_p = hdr_body[0]
        if not ''.join(first_p.itertext()).strip():
            hdr_body.remove(first_p)

    # --- FOOTER: 3-line centered block (left) | Page X of Y (right) | rule line above ---
    footer_table = new_section.footer.add_table(rows=1, cols=2, width=Inches(6.5))
    footer_table.autofit = False
    footer_table.columns[0].width = Inches(5.0)
    footer_table.columns[1].width = Inches(1.5)

    left_cell = footer_table.cell(0, 0)
    right_cell = footer_table.cell(0, 1)

    footer_lines = ["CONFIDENTIAL DOCUMENT", "Copyright © 2026 Opensource.", "All Rights Reserved."]
    for idx, line in enumerate(footer_lines):
        p = left_cell.paragraphs[0] if idx == 0 else left_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_styling(p, line, is_bold=False, font_size=8)

    right_cell.vertical_alignment = 1  # Center
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    apply_styling(right_para, "Page ", is_bold=False, font_size=9)
    add_page_number_field(right_para.add_run())
    apply_styling(right_para, " of ", is_bold=False, font_size=9)
    add_numpages_field(right_para.add_run())

    for cell in (left_cell, right_cell):
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')
        top.set(qn('w:space'), '4')
        top.set(qn('w:color'), '7F7F7F')
        borders.append(top)
        tcPr.append(borders)

    # Same stray-tab-stop cleanup as the header, for the footer's own
    # required empty paragraph.
    for para in new_section.footer.paragraphs:
        if not para.text.strip():
            para.paragraph_format.tab_stops.clear_all()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

    # --- TABLE: VULNERABILITY ASSESSMENT SCAN REPORT ---
    doc.add_heading('VULNERABILITY ASSESSMENT SCAN REPORT', level=1)
    t1 = doc.add_table(rows=3, cols=2)
    t1.style = 'Table Grid'
    t1.autofit = False  # CRITICAL: Prevents Word from resizing columns automatically

    pn = get_project_name(workspace)
    t1_labels = ["Document Title", "Document File Name", "Document Purpose"]

    for i, l in enumerate(t1_labels):
        row = t1.rows[i]

        # Set specific widths for the columns
        row.cells[0].width = Inches(2.0)  # Smaller first column
        row.cells[1].width = Inches(4.5)  # Larger second column

        row.cells[0].text = l
        value_p = row.cells[1].paragraphs[0]
        if l == "Document Purpose":
            add_text_with_project_name(value_p, '', pn, ' Review showing details of each vulnerability and the current status on the ')
            add_text_with_project_name(value_p, '', pn, ' project servers.')
        else:
            add_text_with_project_name(value_p, '', pn, ' VA Review Final Report Q1-2026')

        set_cell_shading(row.cells[0], "1F497D")
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True

        # Extra top/bottom padding so text isn't cramped against the borders
        set_cell_margins(row.cells[0])
        set_cell_margins(row.cells[1])

    doc.add_paragraph("\n")

    # --- TABLE: SECURITY TEAM DETAILS ---
    doc.add_heading('SECURITY TEAM DETAILS', level=1)
    t2 = doc.add_table(rows=4, cols=2)
    t2.style = 'Table Grid'
    t2.autofit = False

    team = [
        ("Full Name", "Certifications"),
        ("Penigalapati Vinaya Kumar", "*Offensive Security Certified Professional (OSCP) (Certificate ID: OS-101-53426)\n*CREST Registered Penetration Tester (CREST CRT) (CREST ID: 1269386536)"),
        ("Domatoti Vinod Kumar", "*Offensive Security Certified Professional (OSCP) (OS-101-025977)\n*Offensive Security Web Assessor (OSWA) Certificate ID: OSWA-28308\n*Offensive Security Wireless Professional (OSWP) License No: OS-BWA-15192\n*Certified Cybersecurity from (ISC)² - (Certification Number: 1154738)"),
        ("Pranatarthi Ravindra Kumar", "*Offensive Security Certified Professional (OSCP)\n*ISACA: Certified Information Security Auditor (CISA)\n*ISACA: Certified Information Security Manager (CISM)")
    ]

    for i, (name, cert) in enumerate(team):
        row = t2.rows[i]

        # Set widths (matching the first table for consistency)
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.5)

        # Extra top/bottom padding so multi-line cells get breathing room
        set_cell_margins(row.cells[0])
        set_cell_margins(row.cells[1])

        row.cells[0].text = name

        # Certifications: one paragraph per line (instead of one cramped block)
        # with spacing after each, so the list reads like the reference report.
        cert_cell = row.cells[1]
        cert_cell.text = ""
        for idx, line in enumerate(cert.split("\n")):
            p = cert_cell.paragraphs[0] if idx == 0 else cert_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

        if i == 0:
            set_cell_shading(row.cells[0], "1F497D")
            set_cell_shading(row.cells[1], "1F497D")
            for r in row.cells[0].paragraphs[0].runs + row.cells[1].paragraphs[0].runs:
                r.font.color.rgb = RGBColor(255, 255, 255)

 #       if i == 2: # Green row for Vinod
 #           set_cell_shading(row.cells[0], "EBF1DE")
 #           set_cell_shading(row.cells[1], "EBF1DE")

  # --- TABLE: DOCUMENT REVIEW ---
    doc.add_heading('DOCUMENT REVIEW', level=1)
    
    # 11 rows: (5 blocks of Name/Date) + 1 row for Submission Date
    t3 = doc.add_table(rows=11, cols=4)
    t3.style = 'Table Grid'
    t3.autofit = False

    # Set column widths to match previous tables
    # Col 0: Label (1.2"), Col 1: Name/Date (2.8"), Col 2: Sig Label (0.8"), Col 3: Sig Box (1.7")
    widths = [Inches(1.2), Inches(2.8), Inches(0.8), Inches(1.7)]
    for row in t3.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    review_data = [
        ("Prepared by\n(Opensource):", "Domatoti Vinod Kumar"),
        ("Reviewed by\n(Opensource):", "Ankur Dayal"),
        ("Reviewed by\n(%s):" % get_customer_name(workspace), ""),
        ("Accepted by\n(ITCD):", ""),
        ("Accepted by\n(System Owner):", "")
    ]

    # Loop through the 5 main blocks (Rows 0-9)
    for i, (label, name) in enumerate(review_data):
        base_row = i * 2
        
        # 1. Labels (Col 0) - Dark Blue
        cell_label = t3.cell(base_row, 0)
        cell_date_label = t3.cell(base_row + 1, 0)
        
        cell_label.text = label
        cell_date_label.text = "Date:"
        
        set_cell_shading(cell_label, "1F497D ")
        set_cell_shading(cell_date_label, "1F497D")
        
        for cell in [cell_label, cell_date_label]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    r.font.size = Pt(8)
                    r.bold = True
            set_cell_margins(cell, top=100, bottom=100)

        # 2. Values (Col 1)
        value_cell = t3.cell(base_row, 1)
        value_cell.text = name
        set_cell_margins(value_cell, top=100, bottom=100)

        # 3. Signature Column (Col 2) - Merged & Dark Blue
        sig_label_cell = t3.cell(base_row, 2).merge(t3.cell(base_row + 1, 2))
        sig_label_cell.text = "Signature:"
        set_cell_shading(sig_label_cell, "1F497D")
        sig_label_cell.vertical_alignment = 1 # Center
        for p in sig_label_cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(8)
                r.bold = True
        set_cell_margins(sig_label_cell, top=100, bottom=100)

        # 4. Signature Box (Col 3) - Merged
        sig_box_cell = t3.cell(base_row, 3).merge(t3.cell(base_row + 1, 3))
        set_cell_margins(sig_box_cell, top=100, bottom=100)

    # --- Last Row: Submission Date ---
    sub_row = t3.rows[10]
    sub_row.cells[0].text = "Submission\nDate:"
    set_cell_shading(sub_row.cells[0], "1F497D")
    for r in sub_row.cells[0].paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(8)
        r.bold = True
    set_cell_margins(sub_row.cells[0], top=100, bottom=100)

    # Merge the remaining cells for the date value
    merged_date_val = sub_row.cells[1].merge(sub_row.cells[2]).merge(sub_row.cells[3])
    merged_date_val.text = datetime.date.today().strftime("%d %b %Y")
    set_cell_margins(merged_date_val, top=100, bottom=100)

    # Force every row (all 5 name/date blocks + the submission-date row) to
    # the same height so the table reads evenly instead of some rows being
    # taller than others. Kept compact (not much taller than one line of
    # 8pt text + margins) rather than the earlier, visibly-too-tall value.
    for row in t3.rows:
        set_row_height(row, 0.26)

    # --- FOOTNOTE NOTE ---
    note_para = doc.add_paragraph()
    note_run = note_para.add_run("Note: This report will be valid for a duration of 30 Days from the report submission.")
    note_run.font.name = 'Calibri'
    note_run.font.size = Pt(10)
    note_run.italic = True


# --- VERSION HISTORY ---
    doc.add_heading('VERSION HISTORY', level=1)

    # 5 rows total (Header + 1 filled row + 3 empty rows)
    v_table = doc.add_table(rows=5, cols=4)
    v_table.style = 'Table Grid'
    v_table.autofit = False

    # Define widths for the 4 columns
    v_widths = [Inches(1.1), Inches(0.8), Inches(1.1), Inches(3.5)]
    for row in v_table.rows:
        for idx, width in enumerate(v_widths):
            row.cells[idx].width = width

    # Header Row
    v_headers = ["Date", "Version", "Author", "Description / Changes"]
    for i, txt in enumerate(v_headers):
        cell = v_table.rows[0].cells[i]
        cell.text = txt
        set_cell_shading(cell, "1F497D") # Dark Blue
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(9)
                r.bold = True
        set_cell_margins(cell, top=100, bottom=100)

    # Initial Release Data (Row 1)
    v_row1 = v_table.rows[1].cells
    v_row1[0].text = "27 Mar 2026"
    v_row1[1].text = "1.0"
    v_row1[2].text = "Opensource"
    v_row1[3].text = "Initial Release"

    # Set font for all data rows to Calibri 9pt + consistent padding
    for row in v_table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(9)
            set_cell_margins(cell, top=100, bottom=100)

    doc.add_paragraph("\n")

    # --- 4. TABLE OF CONTENTS (Page 4) ---
    doc.add_page_break()
    doc.add_heading('Table of Contents', level=1)
    toc_p = doc.add_paragraph()
    run = toc_p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); run._r.append(f1)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'TOC \\o "1-3" \\h \\z \\u'; run._r.append(it)
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate'); run._r.append(f2)
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end'); run._r.append(f3)

    # --- 3. TABLE OF CONTENTS ---
   # doc.add_heading('Table of Contents', level=1)
   # toc_p = doc.add_paragraph()
   # run = toc_p.add_run()
    #fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin'); run._r.append(fld1)
   # instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-3" \\h \\z \\u'; run._r.append(instr)
   #fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate'); run._r.append(fld2)
   # fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end'); run._r.append(fld3)
    doc.add_page_break()
    

    # Sections 1-3
    # --- 2. SECURITY ANALYSIS OVERVIEW (Page 4/5) ---
    h2 = doc.add_heading('1      PURPOSE OF THIS DOCUMENT', level=1)
    style_section_heading(h2, size=14)

    p_purpose = doc.add_paragraph()
    p_purpose.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text_with_customer_name(p_purpose,
        "This document tabulates the results of the vulnerability assessment conducted for ",
        get_customer_name(workspace),
        " for its client ")
    add_text_with_project_name(p_purpose,
        "",
        get_project_name(workspace),
        ". The tests are performed using Nessus Scan for Vulnerability Assessment Scan in Operating System as well as manual validation. ")
    p_purpose.paragraph_format.space_after = Pt(14)
    doc.add_page_break()


# --- 2. SECURITY ANALYSIS OVERVIEW (Page 4/5) ---

    h2 = doc.add_heading('2      SECURITY ANALYSIS OVERVIEW', level=1)
    style_section_heading(h2, size=14)

    add_body_paragraph_with_project_name(doc,
        "Opensource has been engaged to conduct vulnerability assessment scan for Operating System involved in the ",
        get_project_name(workspace),
        " Project. The scope of this Security Test covers the vulnerability assessment scan for Windows servers.")

    # 2.1 TESTING LANDSCAPE
    style_section_heading(doc.add_heading('2.1      TESTING LANDSCAPE', level=2), size=12)

    add_body_paragraph_with_project_name(doc,
        "This security assessment covers the VA scan for Servers in the PROD, UAT, DEV and MGMT environment of ",
        get_project_name(workspace),
        " project. The assessment was carried out with credentialed access to run the scanner to find the vulnerabilities on the mentioned servers.")

    # 2.2 VULNERABILITY ASSESSMENT SCAN
    style_section_heading(doc.add_heading('2.2      VULNERABILITY ASSESSMENT SCAN', level=2), size=12)
    add_body_paragraph(doc,
        "Vulnerability Assessment Scan is a non-intrusive approach that serves to produce a list of vulnerabilities and "
        "the vulnerability status. A combination of automated and manual scan is performed on the servers, with the "
        "objective to identify the vulnerabilities that are present in the environment.")

    # 2.3 VULNERABILITY CATEGORIZATION GUIDELINES
    style_section_heading(doc.add_heading('2.3      VULNERABILITY CATEGORIZATION GUIDELINES', level=2), size=12)
    add_body_paragraph(doc,
        "Vulnerability Severity Ratings are globally based on CVSS (Common Vulnerability Scoring System) ratings, "
        "which are provided by the Nessus Scanner.")

    # --- CVSS RATINGS TABLE ---
    t_cvss = doc.add_table(rows=5, cols=2)
    t_cvss.style = 'Table Grid'
    t_cvss.autofit = False
    # Set widths: Col 0 (1.5"), Col 1 (5.0")
    for row in t_cvss.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.0)

    cvss_data = [
        ("Severity", "Rating"),
        ("CRITICAL", "The plugin's highest vulnerability CVSSv3 score is between 9.0 to 10.0."),
        ("HIGH", "The plugin's highest vulnerability CVSSv3 score is between 7.0 to 8.9."),
        ("MEDIUM", "The plugin's highest vulnerability CVSSv3 score is between 4.0 to 6.9."),
        ("LOW", "The plugin's highest vulnerability CVSSv3 score is between 0.1 to 3.9.")
    ]

    for i, (sev, rat) in enumerate(cvss_data):
        row = t_cvss.rows[i]
        row.cells[0].text = sev
        row.cells[1].text = rat
        if i == 0:
            set_cell_shading(row.cells[0], "1F497D")
            set_cell_shading(row.cells[1], "1F497D")
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs: r.font.color.rgb = RGBColor(255, 255, 255); r.bold = True
                pad_header_cell(cell)
        else:
            for p in row.cells[0].paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for cell in row.cells:
                set_cell_margins(cell, top=90, bottom=90)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Caption for Table
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_styling(caption, "Table-2.3.1: Nessus CVSS Ratings", is_bold=True, font_size=9)

    va_details_p = doc.add_paragraph("VA Scanner Details:", style='Normal')
    va_details_p.runs[0].bold = True
    va_details_p.paragraph_format.space_after = Pt(10)

    # --- VA SCANNER DETAILS TABLE ---
    t_scan = doc.add_table(rows=5, cols=2)
    t_scan.style = 'Table Grid'
    t_scan.autofit = False
    
    scan_details = [
        ("Tool Used", "Nessus Professional"),
        ("Nessus version", "10.8.4"),
        ("Plugin feed version", "202602181602"),
        ("Scan Type", "VA Credential Scan"),
        ("Final Round Scanned", "25 Mar 2026")
    ]

    for i, (label, val) in enumerate(scan_details):
        row = t_scan.rows[i]
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.7)
        row.cells[0].text = label
        row.cells[1].text = val
        set_cell_shading(row.cells[0], "1F497D")
        for r in row.cells[0].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True
        pad_header_cell(row.cells[0])
        set_cell_margins(row.cells[1], top=90, bottom=90)
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Caption for Table
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_styling(caption, "Table-2.3.2: VA Scanner details", is_bold=True, font_size=9)

    doc.add_page_break()
# --- 2.4 VULNERABILITY ASSESSMENT SCAN STATUS ---
    style_section_heading(doc.add_heading('2.4      VULNERABILITY ASSESSMENT SCAN STATUS', level=2), size=12)

    add_body_paragraph(doc,
        "VA Scan status is provided based on the current settings at the time of scan or manual validation. "
        "Please refer the below table for the status used in this VA Scan report.")

    vuln_status_p = doc.add_paragraph("Vulnerability Status:", style='Normal')
    vuln_status_p.runs[0].bold = True
    vuln_status_p.paragraph_format.space_after = Pt(10)

    # Create Table (4 rows, 2 columns)
    t_status = doc.add_table(rows=4, cols=2)
    t_status.style = 'Table Grid'
    t_status.autofit = False
    
    # Set widths (Col 0: 2.0", Col 1: 4.5")
    for row in t_status.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.5)

    # Header Row
    h_row = t_status.rows[0]
    h_row.cells[0].text = "Vulnerability Status"
    h_row.cells[1].text = "Description"
    set_cell_shading(h_row.cells[0], "1F497D")
    set_cell_shading(h_row.cells[1], "1F497D")
    for cell in h_row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True
                r.font.size = Pt(10)
        pad_header_cell(cell)

    # Status Data with specific colors
    status_definitions = [
        ("Open", "The vulnerability status will be marked as Open when a vulnerability is not fixed. The Accenture team need to fix this and inform Opensource Security team for validation. Any Re-Opened vulnerabilities will also be marked as Open.", RGBColor(255, 0, 0)), # Red
        ("Fixed & Closed", "The vulnerability status will be marked as Fixed when it is mitigated by Accenture team, and Opensource Security team has validated and confirmed that the vulnerability is Fixed.", RGBColor(0, 176, 80)), # Green
        ("Seeking Risk Acceptance", "If a vulnerability cannot be fixed due to various reasons or if it is an expected business logic as per business requirements, it will be discussed, and the client should agree to add it to the Risk Register, then the status is marked as \"Seeking Risk Acceptance\". The Infra/Application Team should provide proper justification and supporting evidence for this and any further discussion or follow ups will be done from Risk Register.", RGBColor(0, 176, 240)) # Blue
    ]

    for i, (status, desc, color) in enumerate(status_definitions, 1):
        row = t_status.rows[i]
        
        # Style the Status Column (Column 0) with specific colors
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run0 = p0.add_run(status)
        run0.font.name = 'Calibri'
        run0.font.size = Pt(10)
        run0.font.color.rgb = color
        run0.bold = True

        # Style the Description Column (Column 1)
        p1 = row.cells[1].paragraphs[0]
        run1 = p1.add_run(desc)
        run1.font.name = 'Calibri'
        run1.font.size = Pt(9)

        for cell in row.cells:
            set_cell_margins(cell, top=90, bottom=90)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Table Caption
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_styling(cap_para, "Table-2.4.1: Vulnerability status used in the detailed VA Scan report.", is_bold=True, font_size=9)

    doc.add_page_break()

    # --- 3. VULNERABILITY ASSESSMENT SCAN SUMMARY (section intro) ---
    h3 = doc.add_heading('3      VULNERABILITY ASSESSMENT SCAN SUMMARY', level=1)
    style_section_heading(h3, size=14)

    add_body_paragraph(doc,
        "The following table shows the total vulnerabilities identified during the VA scan and their current status.")

    note3_p = doc.add_paragraph()
    note3_run = apply_styling(note3_p, "Note: This report reflects the security status of the environment as of "
        "the date when the security assessment (including pre-remediation and post-remediation round) was "
        "conducted. Any changes, modifications or alterations to the environment post the aforesaid assessment "
        "are not covered under this report.", font_size=10)
    note3_run.italic = True
    note3_p.paragraph_format.space_after = Pt(14)

    # --- 3.1 VULNERABILITY ASSESSMENT SCAN SUMMARY ---
    style_section_heading(doc.add_heading('3.1      VULNERABILITY ASSESSMENT SCAN SUMMARY', level=2), size=12)

    add_body_paragraph(doc,
        "The following table depicts the summary of total number of total vulnerabilities from the Round-1 & final Round of scan.")

    # 17 columns: S.No(1), Server(1), Crit(3), High(3), Med(3), Low(3), Grand Total(3)
    summary_table = doc.add_table(rows=2, cols=17)
    summary_table.style = 'Table Grid'
    summary_table.autofit = False

    # Adjust widths to fit all 17 columns (Total printable width ~6.5-7.0 inches)
    widths = [0.3, 1.0] + [0.4] * 15 
    for i, w in enumerate(widths):
        for row in summary_table.rows:
            row.cells[i].width = Inches(w)

    def style_header_cell(cell, text, color):
        cell.text = text
        set_cell_shading(cell, color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True
        run.font.size = Pt(8)
        set_cell_margins(cell, top=50, bottom=50, left=40, right=40)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Base Headers
    summary_table.cell(0, 0).merge(summary_table.cell(1, 0))
    style_header_cell(summary_table.cell(0, 0), "S.No", "1F497D")
    summary_table.cell(0, 1).merge(summary_table.cell(1, 1))
    style_header_cell(summary_table.cell(0, 1), "Server Details", "1F497D")

    # Severity Group Headers
    style_header_cell(summary_table.cell(0, 2).merge(summary_table.cell(0, 4)), "Critical", "C00000") 
    style_header_cell(summary_table.cell(0, 5).merge(summary_table.cell(0, 7)), "High", "E46C0A")    
    style_header_cell(summary_table.cell(0, 8).merge(summary_table.cell(0, 10)), "Medium", "FFC000") 
    style_header_cell(summary_table.cell(0, 11).merge(summary_table.cell(0, 13)), "Low", "0070C0")
    # NEW: Grand Total Header
    style_header_cell(summary_table.cell(0, 14).merge(summary_table.cell(0, 16)), "Grand Total", "1F4E78") # Darker Navy Blue

    # Sub-headers (Open, Fixed, Total)
    sub_titles = ["Open", "Fixed", "Total"] * 5
    for i, title in enumerate(sub_titles, 2):
        cell = summary_table.cell(1, i)
        cell.text = title
        set_cell_shading(cell, "365F91")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(255, 255, 255)
        if "Open" in title: run.font.color.rgb = RGBColor(255, 0, 0)
        elif "Fixed" in title: run.font.color.rgb = RGBColor(0, 176, 80)
        set_cell_margins(cell, top=50, bottom=50, left=40, right=40)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # --- DATA CALCULATION ---
    hosts = df['Host'].unique()
    severities = ['Critical', 'High', 'Medium', 'Low']
    
    # Track overall column totals for footer
    col_totals = [0] * 15 # Index 0-14 for the 15 data columns

    for idx, host in enumerate(hosts, 1):
        row_cells = summary_table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = host
        host_df = df[df['Host'] == host]
        
        row_open_sum = 0
        row_fixed_sum = 0
        
        col_idx = 2
        for sev in severities:
            open_cnt = len(host_df[(host_df['Severity'].str.lower() == sev.lower()) & (host_df['Vulnerability Status'] == 'Open')])
            fixed_cnt = len(host_df[(host_df['Severity'].str.lower() == sev.lower()) & (host_df['Vulnerability Status'] == 'Fixed & Closed')])
            
            total_cnt = open_cnt + fixed_cnt
            row_cells[col_idx].text = str(open_cnt)
            row_cells[col_idx+1].text = str(fixed_cnt)
            row_cells[col_idx+2].text = str(total_cnt)
            
            # Update row sums for grand total
            row_open_sum += open_cnt
            row_fixed_sum += fixed_cnt
            
            # Update footer totals
            col_totals[col_idx-2] += open_cnt
            col_totals[col_idx-1] += fixed_cnt
            col_totals[col_idx] += total_cnt
            
            col_idx += 3
        
        # Fill Grand Total for this row
        row_cells[14].text = str(row_open_sum)
        row_cells[15].text = str(row_fixed_sum)
        row_cells[16].text = str(row_open_sum + row_fixed_sum)
        
        # Update grand total footer sums
        col_totals[12] += row_open_sum
        col_totals[13] += row_fixed_sum
        col_totals[14] += (row_open_sum + row_fixed_sum)

    # --- FOOTER ---
    foot = summary_table.add_row().cells
    foot[0].merge(foot[1])
    foot[0].text = "Total Count"
    foot[0].paragraphs[0].alignment = 1
    
    for i in range(15):
        cell = foot[i+2]
        cell.text = str(col_totals[i])
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_styling(cap, "Table-3.1.1: Vulnerability Assessment Scan Summary", is_bold=True, font_size=9)


    # --- Data Processing: Unique Vulnerabilities ---
  # --- VULNERABILITIES SUMMARY ---
    # 1. Bold Title
    p_sum_title = doc.add_paragraph()
    run_sum_title = p_sum_title.add_run("Vulnerabilities Summary:")
    run_sum_title.bold = True
    run_sum_title.font.name = 'Calibri'
    run_sum_title.font.size = Pt(11)
    p_sum_title.paragraph_format.space_after = Pt(10)

    add_text_with_project_name(doc.add_paragraph(),
        "The table below displays the vulnerability names identified in the ",
        get_project_name(workspace),
        " environment scan for Round 1 and final Round, excluding any duplicate findings. The severity and vulnerability status have been updated based on the final Round scan results.")

    # 2. Data Processing (MOVE THIS ABOVE THE TABLE LOGIC)
    unique_findings = df.drop_duplicates(subset=['Name']).copy()
    severity_order = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3}
    unique_findings['sev_rank'] = unique_findings['Severity'].map(severity_order).fillna(4)
    unique_findings = unique_findings.sort_values('sev_rank')

    # This severity-sorted order is the single source of truth for finding
    # numbers (VA_N / 4.1.N). Section 4 below must walk findings in this
    # SAME order and reuse these SAME numbers - otherwise "4.1.17" in this
    # summary table and "4.1.17" in section 4 end up being two completely
    # different vulnerabilities, which makes the statuses look contradictory
    # when they're actually just talking about different findings.
    finding_order = list(unique_findings['Name'])
    name_to_num = {name: idx for idx, name in enumerate(finding_order, 1)}

    # 3. Create Table
    v_sum_table = doc.add_table(rows=1, cols=5)
    v_sum_table.style = 'Table Grid'
    v_sum_table.autofit = False
    v_sum_table.allow_break_across_pages = True
    
    # Force Fixed Layout to prevent margin overflow
    v_sum_table._element.xpath('./w:tblPr/w:tblLayout')[0].set(qn('w:type'), 'fixed')

    # Total width = 6.5"
    v_widths = [Inches(0.6), Inches(0.9), Inches(3.0), Inches(0.9), Inches(1.1)]

    # 4. Header Row
    v_headers = ["S.No", "Vulnerability ID", "Name of Vulnerability", "Severity", "Vulnerability Status"]
    for i, txt in enumerate(v_headers):
        cell = v_sum_table.rows[0].cells[i]
        cell.width = v_widths[i] # Set width
        cell.text = txt
        set_cell_shading(cell, "1F497D")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
        run.bold = True
        pad_header_cell(cell, top=60, bottom=60)

    # 5. Fill Data Rows
    for idx, (index, row_data) in enumerate(unique_findings.iterrows(), 1):
        row = v_sum_table.add_row().cells
        
        # Apply fixed widths and vertical alignment to every row
        for i, w in enumerate(v_widths):
            row[i].width = w
            row[i].vertical_alignment = 1 # 1 = Center

        row[0].text = f"4.1.{idx}"
        row[1].text = f"VA_{idx}"
        row[2].text = str(row_data['Name'])
        
        sev_val = str(row_data['Severity'])
        row[3].text = sev_val
        # Aggregate across every server reported for this finding, not just
        # this one (deduped) row - a finding is only "Fixed & Closed" once
        # ALL of its servers are fixed.
        name_statuses = df.loc[df['Name'] == row_data['Name'], 'Vulnerability Status']
        row[4].text = compute_overall_status(name_statuses)

        # Cell Formatting
        for i in range(5):
            if i != 2: # Center align everything except the name
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for p in row[i].paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(10)

        # Severity Coloring
        sev_p = row[3].paragraphs[0]
        if sev_p.runs:
            sev_run = sev_p.runs[0]
            sev_clean = sev_val.lower()
            if 'critical' in sev_clean or 'high' in sev_clean:
                sev_run.font.color.rgb = RGBColor(255, 0, 0)
            elif 'medium' in sev_clean:
                sev_run.font.color.rgb = RGBColor(228, 108, 10)
            elif 'low' in sev_clean:
                sev_run.font.color.rgb = RGBColor(0, 112, 192)

    # 6. Add Table Caption
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_with_project_name(cap_para, "Table-3.1.2: Vulnerabilities Summary identified in ", get_project_name(workspace), "", is_bold=True, font_size=9)

    doc.add_paragraph("\n")

    # --- 3.2 VA SCAN PROCESS ---
    style_section_heading(doc.add_heading('3.2      VA SCAN PROCESS', level=2), size=12)

    add_body_paragraph(doc, "Here is the VA Scan process in brief.")

    # List items from screenshot - each is (before, after) split around
    # where the customer/client team name goes.
    process_steps = [
        ("Upon completion of the initial round of VA scan on the PROD, UAT, DEV and MGMT environment, "
         "the Opensource security team disseminated both the Nessus HTML and prepared excel reports to the ",
         " team."),

        ("The ",
         " team implemented the necessary fixes and requested that the Opensource security team "
         "conduct a final round of VA Scan on the PROD, UAT, DEV and MGMT environment."),

        ("Upon the completion of the final round of VA scans on the PROD, UAT, DEV and MGMT environment, "
         "the Opensource security team disseminated both the Nessus HTML report and the prepared Excel report to the ",
         " team."),

        ("The Opensource security team validated the vulnerabilities identified in the final round of VA scans, "
         "updated the VA scan report, prepared a Word report, and shared both reports with ",
         " team.")
    ]

    # Native Word numbered-list style instead of a manually-typed "1.    text" prefix
    cust_name = get_customer_name(workspace)
    for before_txt, after_txt in process_steps:
        p = doc.add_paragraph(style='List Number')
        add_text_with_customer_name(p, before_txt, cust_name, after_txt, font_size=10)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()
    # --- SECTION 4: FINDINGS ---

    # --- 4. DETAILED VULNERABILITY ASSESSMENT SCAN REPORT Header ---
    h2 = doc.add_heading('4      DETAILED VULNERABILITY ASSESSMENT SCAN REPORT', level=1)
    style_section_heading(h2, size=14)

    

    # Walk findings in the SAME severity-sorted order used for the
    # Vulnerabilities Summary table above, and reuse the SAME finding
    # numbers (name_to_num), so "4.1.N" / "VA_N" refers to one and only one
    # vulnerability across the whole report instead of drifting between
    # the two tables.
    grouped = df.groupby('Name', sort=False)
    for name in finding_order:
        i = name_to_num[name]
        group = grouped.get_group(name)
        h_finding = doc.add_heading(f'4.1.{i} {name}', level=2)
        for run in h_finding.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black instead of theme blue
        h_finding.paragraph_format.space_before = Pt(14)
        h_finding.paragraph_format.space_after = Pt(6)
        table = doc.add_table(rows=20, cols=3)
        table.style = 'Table Grid'
        unique_group = group.drop_duplicates(subset=['Host', 'Port'])

        # R1-2: Meta Info
        for idx, txt in enumerate(['Vulnerability ID', 'Severity', 'Vulnerability Status']):
            cell = table.rows[0].cells[idx]
            set_cell_background(cell, "1F497D")
            apply_styling(cell.paragraphs[0], txt, is_bold=True, font_size=10, color=RGBColor(255, 255, 255))
            pad_header_cell(cell)

        for idx in range(3):
            set_cell_margins(table.rows[1].cells[idx], top=90, bottom=90)
            table.rows[1].cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        apply_styling(table.rows[1].cells[0].paragraphs[0], f'VA_{i}')
        apply_styling(table.rows[1].cells[1].paragraphs[0], str(group.iloc[0].get('Severity', 'N/A')).upper())
        # A finding only counts as Fixed & Closed once every server it was
        # found on is fixed - one still-open server means the finding is
        # still Open, even if the rest were remediated.
        status_val = compute_overall_status(group['Vulnerability Status'].values)
        apply_styling(table.rows[1].cells[2].paragraphs[0], status_val)

        # R3-4: Vulnerability Name
        format_full_width_header(table, 2, "Vulnerability")
        apply_styling(format_full_width_value(table, 3).paragraphs[0], name, font_size=10)

        # R5-6: Affected Servers
        format_full_width_header(table, 4, "Affected servers")
        affected_cell = format_full_width_value(table, 5)
        status_order = ['Fixed & Closed', 'Open']
        found_statuses = [s for s in status_order if s in unique_group['Vulnerability Status'].values]

        for s_idx, status_type in enumerate(found_statuses):
            sub = unique_group[unique_group['Vulnerability Status'] == status_type]
            p6 = affected_cell.paragraphs[0] if s_idx == 0 else affected_cell.add_paragraph()
            apply_styling(p6, "Current Status: ", is_bold=True, font_size=10)
            apply_styling(p6, f"{status_type}", font_size=10)
            ips = format_unique_servers(sub.itertuples())
            apply_styling(p6, "\nServers: ", is_bold=True, font_size=10)
            apply_styling(p6, f"{ips}", font_size=10)
            p6.paragraph_format.space_after = Pt(8)

        # R7-8: Synopsis
        format_full_width_header(table, 6, "Issue Synopsis")
        synopsis_val = group['Synopsis'].iloc[0] if not group['Synopsis'].empty else "N/A"
        add_cleaned_text_block(format_full_width_value(table, 7), synopsis_val, font_size=10)

        # R9-10: Findings (With Plugin Output Fix & Italic Note) - each sample gets its own paragraph
        format_full_width_header(table, 8, "Findings")
        findings_cell = format_full_width_value(table, 9)
        findings_df = group.dropna(subset=['Plugin Output']).drop_duplicates(subset=['Plugin Output'])[:3]

        for idx, (_, row_data) in enumerate(findings_df.iterrows(), 1):
            p10 = findings_cell.paragraphs[0] if idx == 1 else findings_cell.add_paragraph()
            apply_styling(p10, f"Sample Finding Output-{idx}: ", is_bold=True, font_size=10)
            apply_styling(p10, f"{row_data['Host']}(Port:{row_data['Port']})\n", font_size=10)
            apply_styling(p10, f"{row_data['Plugin Output']}", font_size=10)
            p10.paragraph_format.space_after = Pt(8)

        # --- ITALIC NOTE ---
        note_text = "Note: More than 3 outputs present, hence restricted to 3 output samples. For more details, refer to HTML"
        note_p = findings_cell.paragraphs[0] if findings_df.empty else findings_cell.add_paragraph()
        note_run = apply_styling(note_p, note_text, is_bold=False, font_size=10)
        note_run.italic = True

       # R11-12: Description
        format_full_width_header(table, 10, "Description")
        # iloc[0] used for common data; displays empty if nan
        desc_val = str(group['Description'].iloc[0])
        add_cleaned_text_block(format_full_width_value(table, 11), "" if desc_val.lower() == 'nan' else desc_val, font_size=10, fallback="")

        # R13-14: Recommendation
        format_full_width_header(table, 12, "Recommendation")
        sol_val = str(group['Solution'].iloc[0])
        add_cleaned_text_block(format_full_width_value(table, 13), "" if sol_val.lower() == 'nan' else sol_val, font_size=10, fallback="")

        # R15 - R20: Comments & Artifacts
        comment_sets = [
            (14, 15, "Customer comments", "Customer Comment", ["Customer Comments", "Customer Comment"]),
            (16, 17, "Opensource comment", "Opensource Comment", ["Opensource Comments", "Opensource Comment"]),
            (18, 19, "Artifacts", "Artifact", ["Artefacts", "artifacts", "Artifact", "Artefact"])
        ]

        # Clean the group for logic but keep values empty strings
        unique_group = group.fillna('') 

        for h_idx, v_idx, h_title, prefix, possible_cols in comment_sets:
            format_full_width_header(table, h_idx, h_title)
            val_cell = format_full_width_value(table, v_idx)

            # Find column case-insensitively
            actual_col = next((col for col in group.columns if col.strip().lower() in [c.lower() for c in possible_cols]), None)

            if actual_col:
                # Only grab values that are actually filled
                unique_vals = [v for v in unique_group[actual_col].unique() if str(v).strip() != '']

                for v_sub_idx, val in enumerate(unique_vals):
                    sub_srv = unique_group[unique_group[actual_col] == val]
                    # One paragraph per entry, with spacing after, instead of one crammed block
                    p_val = val_cell.paragraphs[0] if v_sub_idx == 0 else val_cell.add_paragraph()

                    # Add prefix in bold, then the value
                    apply_styling(p_val, f"{prefix}: ", is_bold=True, font_size=10)
                    apply_styling(p_val, f"{val}", font_size=10)

                    # Add Servers mapping
                    ips = format_unique_servers(r for r in sub_srv.itertuples() if hasattr(r, 'Host'))
                    apply_styling(p_val, "\nServers: ", is_bold=True, font_size=10)
                    apply_styling(p_val, f"{ips}", font_size=10)
                    p_val.paragraph_format.space_after = Pt(8)
            else:
                # If column doesn't exist, val_cell remains blank
                val_cell.paragraphs[0].text = ""

        doc.add_paragraph()


    doc.add_page_break()

    # --- 5	APPENDIX -A ---
    # --- 5 APPENDIX -A ---
    doc.add_page_break()
    h5 = doc.add_heading('APPENDIX -A', level=1)
    style_section_heading(h5, size=14)

    add_body_paragraph(doc, "The following table contains the asset list and Nessus VA scan reports.")

    # Create Table: 5 rows (Header + 4 items), 3 columns
    app_table = doc.add_table(rows=5, cols=3)
    app_table.style = 'Table Grid'
    app_table.autofit = False
    
    # Force Fixed Layout to keep the table strictly within page margins
    app_table._element.xpath('./w:tblPr/w:tblLayout')[0].set(qn('w:type'), 'fixed')

    # Optimized Widths: S.No (0.4"), Description (3.05"), Artefacts (3.05")
    v_widths = [Inches(0.4), Inches(3.05), Inches(3.05)]

    # Header Row Styling
    headers = ["S.No", "Description", "Artefacts"]
    for i, txt in enumerate(headers):
        cell = app_table.rows[0].cells[i]
        cell.width = v_widths[i]
        cell.text = txt
        set_cell_shading(cell, "1F497D") 
        p = cell.paragraphs[0]
        p.alignment = 1 # Center
        run = p.runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True
        run.font.size = Pt(10)
        pad_header_cell(cell)

    # Content Data (Rows 2 & 3 are now just empty placeholders)
    content_rows = [
        ("1", "List of Servers and the IP address details"),
        ("2", "VA Initial Round Scan Report"),
        ("3", "VA Final Round Scan Report"),
        ("4", "VA Scan results in Excel format")
    ]

    for idx, (sno, desc) in enumerate(content_rows, 1):
        row = app_table.rows[idx].cells
        
        # Apply fixed widths to every row cell
        for i, w in enumerate(v_widths):
            row[i].width = w
            row[i].vertical_alignment = 1 # Center everything vertically
            
        row[0].text = sno
        row[1].text = desc
        
        # Artefacts cell (Col 2) is left blank for manual insertion
        # We ensure it has a paragraph so you can click into it easily
        p_art = row[2].paragraphs[0]
        p_art.alignment = 1 
        p_art.text = "" 

    # Final Formatting Loop (Setting Row Heights/Spacing)
    for row in app_table.rows[1:]:
        # Set a minimum row height to make them look like "boxes" for icons
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "800") # approx 0.55 inches
        trHeight.set(qn('w:hRule'), "atLeast")
        trPr.append(trHeight)

        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(10)
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(10)

    # Table Caption
    cap = doc.add_paragraph()
    cap.alignment = 1
    apply_styling(cap, "Table-5.A: List of Artefacts", is_bold=True, font_size=9)


    # Force every table in the document flush to the left margin (same
    # starting edge as the headings/paragraphs above them).
    for tbl in doc.tables:
        zero_table_indent(tbl)

    # Generate a unique filename using a timestamp
    timestamp = time.strftime("%Y%m%d-%H%M%S")
    save_filename = f"VA_Report_{timestamp}.docx"

    save_path = os.path.join(settings.MEDIA_ROOT, save_filename)
    doc.save(save_path)
    
    return save_filename